import streamlit as st
import pandas as pd
from io import BytesIO

# 워드 라이브러리
try:
    from docx import Document
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False

# 페이지 설정
st.set_page_config(
    page_title="헌금 이름 정렬",
    page_icon="📋",
    layout="wide"
)

# 스타일
st.markdown("""
<style>
.main {
    background: linear-gradient(135deg, #f8f6f2 0%, #fdfcf9 100%);
}

h1 {
    color: #4b3f35;
}

section[data-testid="stSidebar"] {
    background: #f3efe8;
}

.stDownloadButton button {
    border-radius: 10px;
}
</style>
""", unsafe_allow_html=True)

# 사이드바
st.sidebar.image("logo.png", width=100)
st.sidebar.markdown("---")

user_name = st.sidebar.text_input(
    "사용자",
    placeholder="대망교회"
)

st.sidebar.info(
    "엑셀 업로드 후 자동 정렬됩니다"
)

# 메인
st.title("📋 헌금 이름 정렬 프로그램")

uploaded_file = st.file_uploader(
    "엑셀 파일 업로드",
    type=["xlsx", "xls"]
)

if uploaded_file is not None:

    df = pd.read_excel(uploaded_file)

    # 완전히 빈 열 제거
    df = df.dropna(axis=1, how="all")

    # Unnamed 열 제거
    df = df.loc[
        :,
        ~df.columns.astype(str).str.contains("^Unnamed")
    ]

    output_text = ""

    # 고정 헌금 순서
    ORDER = [
        "십일조",
        "감사헌금",
        "선교헌금",
        "소원헌금",
        "건축헌금",
        "한국교회살리기헌금",
        "해외선교헌금",
        "생일감사헌금",
        "심방감사헌금",
        "작정헌금",
        "일천번제"
    ]

    existing_cols = [
        str(col).strip()
        for col in df.columns
    ]

    before_extra = [
        col
        for col in ORDER[:4]
        if col in existing_cols
    ]

    extra_cols = [
        col
        for col in existing_cols
        if col not in ORDER
    ]

    after_extra = [
        col
        for col in ORDER[4:]
        if col in existing_cols
    ]

    sorted_columns = (
        before_extra +
        extra_cols +
        after_extra
    )

    for column in sorted_columns:

        title = str(column).strip()

        if title == "":
            continue

        col_data = (
            df[column]
            .dropna()
            .astype(str)
            .str.strip()
        )

        col_data = [
            x
            for x in col_data
            if (
                x
                and x.lower() not in [
                    "nan",
                    "none",
                    "noname"
                ]
            )
        ]

        if len(col_data) == 0:
            continue

        sorted_col = sorted(col_data)

        output_text += f"[{title}]\n"

        for i in range(0, len(sorted_col), 17):
            line = " ".join(
                sorted_col[i:i + 17]
            )
            output_text += line + "\n"

        output_text += "\n"

    if user_name:
        st.success(
            f"{user_name}님 결과입니다"
        )

    st.text_area(
        "결과",
        output_text,
        height=400
    )

    txt_data = output_text.encode(
        "utf-8-sig"
    )

    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            label="📄 텍스트 다운로드 (.txt)",
            data=txt_data,
            file_name="헌금정리결과.txt",
            mime="text/plain",
            key="txt_download"
        )

    if DOCX_AVAILABLE:

        doc = Document()

        if user_name:
            doc.add_heading(
                f"{user_name}님 결과",
                0
            )
        else:
            doc.add_heading(
                "헌금 이름 정렬 결과",
                0
            )

        for line in output_text.split("\n"):
            doc.add_paragraph(line)

        buffer = BytesIO()

        doc.save(buffer)

        buffer.seek(0)

        with col2:
            st.download_button(
                label="📘 워드 다운로드 (.docx)",
                data=buffer,
                file_name="헌금정리결과.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="docx_download"
            )

    else:
        st.warning(
            "⚠️ 워드 다운로드를 사용하려면 python-docx 설치 필요"
        )